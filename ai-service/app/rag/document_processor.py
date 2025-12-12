"""
Processador de documentos para extração de texto e chunking.
Suporta PDF, DOCX, TXT e Markdown.
"""
import io
import re
from typing import List, Optional
from dataclasses import dataclass

import structlog

logger = structlog.get_logger()


@dataclass
class TextChunk:
    """Representa um chunk de texto extraído."""
    text: str
    start_index: int
    end_index: int
    metadata: dict


@dataclass
class ProcessedDocument:
    """Resultado do processamento de um documento."""
    filename: str
    file_type: str
    raw_text: str
    chunks: List[TextChunk]
    metadata: dict


class DocumentProcessor:
    """
    Processa documentos e extrai texto em chunks.
    
    Suporta:
    - PDF (via PyPDF2 ou pdfplumber)
    - DOCX (via python-docx)
    - TXT/MD (direto)
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Args:
            chunk_size: Tamanho máximo de cada chunk em caracteres
            chunk_overlap: Sobreposição entre chunks para manter contexto
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    async def process_file(
        self,
        content: bytes,
        filename: str,
        file_type: str
    ) -> ProcessedDocument:
        """
        Processa um arquivo e retorna chunks de texto.
        
        Args:
            content: Conteúdo do arquivo em bytes
            filename: Nome do arquivo
            file_type: Extensão do arquivo (.pdf, .docx, etc.)
        
        Returns:
            ProcessedDocument com texto e chunks
        """
        # Extrai texto baseado no tipo
        if file_type in ['.pdf']:
            raw_text = await self._extract_pdf(content)
        elif file_type in ['.docx', '.doc']:
            raw_text = await self._extract_docx(content)
        elif file_type in ['.txt', '.md']:
            raw_text = content.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Tipo de arquivo não suportado: {file_type}")
        
        # Limpa e normaliza texto
        raw_text = self._clean_text(raw_text)
        
        # Divide em chunks
        chunks = self._create_chunks(raw_text)
        
        return ProcessedDocument(
            filename=filename,
            file_type=file_type,
            raw_text=raw_text,
            chunks=chunks,
            metadata={
                'original_length': len(raw_text),
                'chunk_count': len(chunks),
            }
        )
    
    async def _extract_pdf(self, content: bytes) -> str:
        """Extrai texto de PDF."""
        try:
            # Tenta PyPDF2 primeiro
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(content))
                text_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                return '\n\n'.join(text_parts)
            except ImportError:
                pass
            
            # Tenta pdfplumber como fallback
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    return '\n\n'.join(text_parts)
            except ImportError:
                pass
            
            # Se nenhum disponível, loga erro
            logger.warning("pdf_extraction_unavailable", 
                message="Instale PyPDF2 ou pdfplumber para suporte a PDF")
            return ""
            
        except Exception as e:
            logger.error("pdf_extraction_error", error=str(e))
            return ""
    
    async def _extract_docx(self, content: bytes) -> str:
        """Extrai texto de DOCX."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Também extrai de tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(
                        cell.text.strip() 
                        for cell in row.cells 
                        if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)
            
            return '\n\n'.join(text_parts)
            
        except ImportError:
            logger.warning("docx_extraction_unavailable",
                message="Instale python-docx para suporte a DOCX")
            return ""
        except Exception as e:
            logger.error("docx_extraction_error", error=str(e))
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Limpa e normaliza texto."""
        if not text:
            return ""
        
        # Remove múltiplas quebras de linha
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove espaços múltiplos
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove caracteres de controle (exceto newlines e tabs)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        
        # Normaliza aspas
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        return text.strip()
    
    def _create_chunks(self, text: str) -> List[TextChunk]:
        """
        Divide texto em chunks com sobreposição.
        
        Tenta dividir em limites naturais (parágrafos, sentenças).
        """
        if not text:
            return []
        
        chunks = []
        
        # Se texto é menor que chunk_size, retorna como único chunk
        if len(text) <= self.chunk_size:
            return [TextChunk(
                text=text,
                start_index=0,
                end_index=len(text),
                metadata={}
            )]
        
        # Divide em parágrafos primeiro
        paragraphs = text.split('\n\n')
        
        current_chunk = ""
        current_start = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Se adicionar este parágrafo excede o limite
            if len(current_chunk) + len(para) + 2 > self.chunk_size:
                if current_chunk:
                    # Salva chunk atual
                    chunks.append(TextChunk(
                        text=current_chunk.strip(),
                        start_index=current_start,
                        end_index=current_start + len(current_chunk),
                        metadata={}
                    ))
                    
                    # Novo chunk com overlap
                    overlap_text = self._get_overlap(current_chunk)
                    current_start = current_start + len(current_chunk) - len(overlap_text)
                    current_chunk = overlap_text + "\n\n" + para if overlap_text else para
                else:
                    # Parágrafo muito grande, divide por sentenças
                    sentences = self._split_into_sentences(para)
                    for sentence in sentences:
                        if len(current_chunk) + len(sentence) + 1 > self.chunk_size:
                            if current_chunk:
                                chunks.append(TextChunk(
                                    text=current_chunk.strip(),
                                    start_index=current_start,
                                    end_index=current_start + len(current_chunk),
                                    metadata={}
                                ))
                                overlap_text = self._get_overlap(current_chunk)
                                current_start = current_start + len(current_chunk) - len(overlap_text)
                                current_chunk = overlap_text + " " + sentence if overlap_text else sentence
                            else:
                                # Sentença muito grande, divide por palavras
                                current_chunk = sentence[:self.chunk_size]
                        else:
                            current_chunk += (" " if current_chunk else "") + sentence
            else:
                current_chunk += ("\n\n" if current_chunk else "") + para
        
        # Adiciona último chunk
        if current_chunk.strip():
            chunks.append(TextChunk(
                text=current_chunk.strip(),
                start_index=current_start,
                end_index=current_start + len(current_chunk),
                metadata={}
            ))
        
        return chunks
    
    def _get_overlap(self, text: str) -> str:
        """Retorna texto para overlap entre chunks."""
        if len(text) <= self.chunk_overlap:
            return text
        
        # Tenta pegar últimas sentenças completas
        overlap_text = text[-self.chunk_overlap:]
        
        # Encontra início de sentença
        sentence_start = overlap_text.find('. ')
        if sentence_start > 0 and sentence_start < len(overlap_text) - 10:
            overlap_text = overlap_text[sentence_start + 2:]
        
        return overlap_text.strip()
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Divide texto em sentenças."""
        # Regex simples para dividir sentenças
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]

